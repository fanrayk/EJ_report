import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, InlineImage
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
import traceback

# ============================================================
# 1. 表格樣式設定 (維持不變：高度3.5、粗體、地址靠左)
# ============================================================

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def create_table_structure(doc, table_type='測量照片'):
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    def format_cell(cell, text="", font_size=None, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_cell_border(cell)
        for run in p.runs:
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            if font_size: run.font.size = font_size
            run.font.bold = bold

    # 標題
    row1 = table.rows[0].cells
    row1[0].merge(row1[3])
    format_cell(row1[0], '欣中天然氣(股)公司 測量作業項目照片', font_size=Pt(18), bold=False)

    # 資訊
    table.rows[1].cells[0].text = '工程案號'
    table.rows[1].cells[1].text = '{{ project_number }}'
    table.rows[1].cells[2].text = '申請書編號'
    table.rows[1].cells[3].text = '{{ application_number }}'
    
    table.rows[2].cells[0].text = '施工地址'
    table.rows[2].cells[1].text = '{{ construction_address }}'
    table.rows[2].cells[2].text = '承攬商'
    table.rows[2].cells[3].text = '庭安科技'

    for r in range(1, 3):
        for c in range(4):
            target_align = WD_ALIGN_PARAGRAPH.CENTER
            if r == 2 and c == 1: 
                target_align = WD_ALIGN_PARAGRAPH.LEFT
            format_cell(table.rows[r].cells[c], table.rows[r].cells[c].text, bold=True, align=target_align)

    format_cell(table.rows[3].cells[0].merge(table.rows[3].cells[3]), table_type, bold=True)

    if table_type == '測量照片':
        for i, row_idx in enumerate([4, 5, 6]):
            row = table.rows[row_idx]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(2.4)
            row.cells[0].merge(row.cells[1])
            row.cells[2].merge(row.cells[3])
            format_cell(row.cells[0], f'{{{{ photo_{i*2+1} }}}}', bold=True)
            format_cell(row.cells[2], f'{{{{ photo_{i*2+2} }}}}', bold=True)
    else:
        for row_idx, tag in [(4, '{{ location_map }}'), (6, '{{ system_screenshot }}')]:
            row = table.rows[row_idx]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(3.5)
            row.cells[0].merge(row.cells[3])
            format_cell(row.cells[0], tag, bold=True)
        format_cell(table.rows[5].cells[0].merge(table.rows[5].cells[3]), '道挖系統上傳完成截圖', bold=True)

    return table

# ============================================================
# 2. 輔助功能：在特定資料夾內找 Excel
# ============================================================

def find_excel_in_folder(target_folder):
    """在指定的資料夾(及其子資料夾)內尋找 Excel"""
    try:
        # 遞迴搜尋 .xlsx 和 .csv
        files = list(target_folder.rglob('*.xlsx')) + list(target_folder.rglob('*.csv'))
        # 排除暫存檔
        valid_files = [f for f in files if not f.name.startswith('~$')]
        
        if valid_files:
            return valid_files[0] # 回傳找到的第一個
    except Exception as e:
        print(f"[DEBUG] 搜尋 Excel 時發生錯誤: {e}")
    return None

# ============================================================
# 3. 單一案場處理邏輯 (獨立載入 Excel)
# ============================================================

def process_single_project(project_dir, template_path):
    print(f"\n[DEBUG] >>> 進入資料夾: {project_dir.name}")
    
    # STEP 1: 在「這個資料夾」裡面找 Excel
    excel_file = find_excel_in_folder(project_dir)
    
    if not excel_file:
        print(f"[DEBUG] ❌ 跳過: 在 {project_dir.name} 裡面找不到 Excel 檔")
        return False
    
    print(f"[DEBUG] 📄 使用 Excel: {excel_file.name}")
    
    # STEP 2: 讀取 Excel
    try:
        if excel_file.suffix == '.csv':
            df = pd.read_csv(excel_file)
        else:
            df = pd.read_excel(excel_file)
    except Exception as e:
        print(f"[DEBUG] ❌ Excel 讀取失敗: {e}")
        return False

    if df.empty:
        print(f"[DEBUG] ❌ Excel 是空的")
        return False

    # 強制轉字串
    df['工程案號'] = df['工程案號'].astype(str).str.strip()
    
    # STEP 3: 決定要用哪一筆資料
    # 邏輯：如果 Excel 只有一筆資料，就直接用那一筆 (最穩)
    # 如果有多筆，嘗試用資料夾名稱匹配
    
    context = {}
    final_project_id = ""
    
    folder_name = project_dir.name
    match_row = df[df['工程案號'] == folder_name]
    
    if len(df) == 1:
        # 單筆資料模式 (適用於搶修/568這種)
        data = df.iloc[0]
        final_project_id = str(data['工程案號'])
        print(f"[DEBUG] 📌 Excel 僅有一筆資料，鎖定案號: {final_project_id}")
    elif not match_row.empty:
        # 匹配成功
        data = match_row.iloc[0]
        final_project_id = str(data['工程案號'])
        print(f"[DEBUG] 📌 資料夾名稱匹配成功，案號: {final_project_id}")
    else:
        # 多筆資料但沒匹配到，預設取第一筆並警告
        data = df.iloc[0]
        final_project_id = str(data['工程案號'])
        print(f"[DEBUG] ⚠️ 無法匹配，預設使用 Excel 第一筆案號: {final_project_id}")

    context = {
        'project_number': final_project_id,
        'application_number': str(data['申請書編號']),
        'construction_address': str(data['施工地址'])
    }

    # STEP 4: 尋找照片
    # 優先找: project_dir / final_project_id / 測量照 (例如 搶修/568/測量照)
    # 其次找: project_dir / 測量照 (例如 06案/測量照)
    
    photo_root = project_dir
    sub_folder_with_id = project_dir / final_project_id
    
    if sub_folder_with_id.exists() and sub_folder_with_id.is_dir():
        photo_root = sub_folder_with_id
        
    print(f"[DEBUG] 📂 照片搜尋根目錄: {photo_root}")

    # STEP 5: 載入範本並填充
    tpl = DocxTemplate(template_path)

    # 測量照
    photo_dir = photo_root / '測量照'
    imgs = sorted(list(photo_dir.glob('*.jpg')) + list(photo_dir.glob('*.png'))) if photo_dir.exists() else []
    print(f"[DEBUG] 📸 找到 {len(imgs)} 張測量照")

    for i in range(1, 7):
        context[f'photo_{i}'] = InlineImage(tpl, str(imgs[i-1]), width=Inches(3.0)) if (i-1) < len(imgs) else ""

    # 其他圖片
    def get_single_img(sub, width):
        d = photo_root / sub
        f = list(d.glob('*.*')) if d.exists() else []
        if f: print(f"[DEBUG] 🖼️  找到 {sub}")
        return InlineImage(tpl, str(f[0]), width=Inches(width)) if f else ""

    context['location_map'] = get_single_img('點位圖', 6.0)
    context['system_screenshot'] = get_single_img('道管截圖', 6.0)

    # STEP 6: 存檔 (檔名使用 Excel 裡的案號)
    tpl.render(context)
    output_filename = f"{final_project_id}_報告書.docx"
    output_path = project_dir / output_filename
    tpl.save(output_path)
    
    print(f"[DEBUG] ✅ 成功產出: {output_filename}")
    return True

# ============================================================
# 4. 主程式流程
# ============================================================

def main_process():
    print("="*50)
    print(">>> 程式啟動 (獨立 Excel 讀取版)")
    print("="*50)
    
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    try:
        # STEP 1: 選擇上層資料夾
        print(">>> 請選擇「上層資料夾」 (包含多個案場資料夾)...")
        root_folder_path = filedialog.askdirectory(title="請選擇上層資料夾", parent=root)
        
        if not root_folder_path:
            return
        
        root_path = Path(root_folder_path)
        print(f"[DEBUG] 根目錄: {root_path}")

        # STEP 2: 準備 Word 範本
        template_name = 'report_template.docx'
        if not os.path.exists(template_name):
            print("[DEBUG] 建立 Word 範本...")
            doc = Document()
            create_table_structure(doc, '測量照片')
            doc.add_page_break()
            create_table_structure(doc, '點位圖')
            doc.save(template_name)

        # STEP 3: 掃描所有子資料夾
        # 這裡不先預判是不是案場，而是進去每個資料夾看有沒有 Excel
        subfolders = [f for f in root_path.iterdir() if f.is_dir()]
        
        if not subfolders:
            messagebox.showwarning("提示", "選擇的資料夾內沒有任何子資料夾")
            return

        print(f"\n[DEBUG] 掃描到 {len(subfolders)} 個子資料夾，開始逐一檢查...\n")
        
        success = 0
        failed = 0
        skipped = 0
        
        for folder in subfolders:
            try:
                # 嘗試處理每個子資料夾
                if process_single_project(folder, template_name):
                    success += 1
                else:
                    skipped += 1
            except Exception as e:
                print(f"[DEBUG] ❌ 處理 {folder.name} 時發生錯誤: {e}")
                traceback.print_exc()
                failed += 1

        # STEP 4: 結束
        print("="*50)
        msg = f"作業結束！\n\n成功生成: {success} 份\n跳過/無Excel: {skipped} 份\n錯誤: {failed} 份"
        print(msg)
        messagebox.showinfo("完成", msg)

    except Exception as e:
        print(f"❌ 嚴重錯誤: {e}")
        traceback.print_exc()
        messagebox.showerror("錯誤", f"發生錯誤: {str(e)}")


if __name__ == '__main__':
    main_process()